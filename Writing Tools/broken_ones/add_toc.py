from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

def create_element(name):
    """Create an OxmlElement with the given name."""
    return OxmlElement(name)

def create_attribute(element, name, value):
    """Create an attribute for the given element."""
    element.set(qn(name), value)

def add_table_of_contents(doc):
    """Add a table of contents to the document after the title page."""
    # Find the first section break (after title page)
    # Insert TOC at the beginning of the document (after title page)
    
    # Add heading for TOC
    toc_heading = doc.add_paragraph("Table of Contents")
    toc_heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
    toc_heading.style = 'Heading 1'
    
    # Add the TOC field
    paragraph = doc.add_paragraph()
    run = paragraph.add_run()
    
    fldChar1 = create_element('w:fldChar')
    create_attribute(fldChar1, 'w:fldCharType', 'begin')
    
    instrText = create_element('w:instrText')
    create_attribute(instrText, 'xml:space', 'preserve')
    instrText.text = 'TOC \\o "1-3" \\h \\z \\u'  # Include heading levels 1-3
    
    fldChar2 = create_element('w:fldChar')
    create_attribute(fldChar2, 'w:fldCharType', 'separate')
    
    fldChar3 = create_element('w:fldChar')
    create_attribute(fldChar3, 'w:fldCharType', 'end')
    
    run._element.append(fldChar1)
    run._element.append(instrText)
    run._element.append(fldChar2)
    run._element.append(fldChar3)
    
    # Add page break after TOC
    doc.add_page_break()

def main():
    # Path to the existing document
    docx_path = r"C:\Users\roger\OneDrive\Desktop\Gauntlet_Projects\PROJECT-REVIEWS\()_SELF-ANALYSIS\AI Engineering Thesis\combined_thesis.docx"
    
    # Path for the new document with TOC
    output_path = r"C:\Users\roger\OneDrive\Desktop\Gauntlet_Projects\PROJECT-REVIEWS\()_SELF-ANALYSIS\AI Engineering Thesis\combined_thesis_with_toc.docx"
    
    # Load the existing document
    doc = Document(docx_path)
    
    # Create a new document for the output
    new_doc = Document()
    
    # Copy the title page (first section) to the new document
    for i, paragraph in enumerate(doc.paragraphs):
        if i < 10:  # Assuming the title page has fewer than 10 paragraphs
            p = new_doc.add_paragraph()
            for run in paragraph.runs:
                new_run = p.add_run(run.text)
                new_run.bold = run.bold
                new_run.italic = run.italic
                new_run.font.name = run.font.name
                if run.font.size:
                    new_run.font.size = run.font.size
            p.alignment = paragraph.alignment
            p.style = paragraph.style
    
    # Add page break after title page
    new_doc.add_page_break()
    
    # Add table of contents
    add_table_of_contents(new_doc)
    
    # Copy the rest of the document
    for i, paragraph in enumerate(doc.paragraphs):
        if i >= 10:  # Skip the title page paragraphs we already copied
            p = new_doc.add_paragraph()
            for run in paragraph.runs:
                new_run = p.add_run(run.text)
                new_run.bold = run.bold
                new_run.italic = run.italic
                new_run.font.name = run.font.name
                if run.font.size:
                    new_run.font.size = run.font.size
            p.alignment = paragraph.alignment
            p.style = paragraph.style
    
    # Copy tables
    for table in doc.tables:
        new_table = new_doc.add_table(rows=len(table.rows), cols=len(table.columns))
        new_table.style = table.style
        
        # Copy cell contents
        for i, row in enumerate(table.rows):
            for j, cell in enumerate(row.cells):
                for paragraph in cell.paragraphs:
                    p = new_table.cell(i, j).paragraphs[0]
                    for run in paragraph.runs:
                        new_run = p.add_run(run.text)
                        new_run.bold = run.bold
                        new_run.italic = run.italic
    
    # Save the new document
    new_doc.save(output_path)
    print(f"Document with table of contents saved as {output_path}")

if __name__ == "__main__":
    main() 