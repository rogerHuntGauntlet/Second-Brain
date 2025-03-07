import os
import re
from pathlib import Path
import markdown
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from bs4 import BeautifulSoup
import docx.oxml.ns as ns
from datetime import datetime
import argparse

def extract_section_number(filename):
    """Extract the section number from the filename."""
    match = re.match(r'(\d+)_', filename)
    if match:
        return int(match.group(1))
    return float('inf')  # For files without section numbers

def create_element(name):
    """Create an OxmlElement with the given name."""
    return OxmlElement(name)

def create_attribute(element, name, value):
    """Create an attribute for the given element."""
    element.set(qn(name), value)

def add_page_number(paragraph):
    """Add page numbers to the document."""
    # Create run for the page number field
    run = paragraph.add_run()
    
    # Create the page number field
    fldChar1 = create_element('w:fldChar')
    create_attribute(fldChar1, 'w:fldCharType', 'begin')
    
    instrText = create_element('w:instrText')
    create_attribute(instrText, 'xml:space', 'preserve')
    instrText.text = "PAGE"
    
    fldChar2 = create_element('w:fldChar')
    create_attribute(fldChar2, 'w:fldCharType', 'end')
    
    # Add field elements to the run
    run._element.append(fldChar1)
    run._element.append(instrText)
    run._element.append(fldChar2)

def setup_styles(doc):
    """Set up document styles for consistent formatting."""
    # Default font and size for the document
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Times New Roman'
    font.size = Pt(12)
    style.paragraph_format.line_spacing_rule = WD_LINE_SPACING.DOUBLE
    style.paragraph_format.space_after = Pt(12)
    
    # Heading styles
    for i in range(1, 7):
        style_name = f'Heading {i}'
        if style_name in doc.styles:
            style = doc.styles[style_name]
            font = style.font
            font.name = 'Times New Roman'
            font.bold = True
            
            if i == 1:
                font.size = Pt(16)
            elif i == 2:
                font.size = Pt(14)
            else:
                font.size = Pt(12)
                
            style.paragraph_format.space_before = Pt(18)
            style.paragraph_format.space_after = Pt(12)
    
    # List styles
    for style_name in ['List Bullet', 'List Number']:
        if style_name in doc.styles:
            style = doc.styles[style_name]
            font = style.font
            font.name = 'Times New Roman'
            font.size = Pt(12)
    
    # Quote style
    if 'Quote' in doc.styles:
        style = doc.styles['Quote']
        font = style.font
        font.name = 'Times New Roman'
        font.italic = True
        font.size = Pt(12)
        style.paragraph_format.left_indent = Inches(0.5)
        style.paragraph_format.right_indent = Inches(0.5)
    
    # Create a style for code blocks
    if 'Code Block' not in doc.styles:
        style = doc.styles.add_style('Code Block', WD_STYLE_TYPE.PARAGRAPH)
        font = style.font
        font.name = 'Courier New'
        font.size = Pt(10)
        style.paragraph_format.left_indent = Inches(0.5)
        style.paragraph_format.space_before = Pt(12)
        style.paragraph_format.space_after = Pt(12)
        style.paragraph_format.line_spacing = 1.0

def apply_heading_style(paragraph, level):
    """Apply heading style based on level."""
    if level == 1:
        paragraph.style = 'Heading 1'
    elif level == 2:
        paragraph.style = 'Heading 2'
    elif level == 3:
        paragraph.style = 'Heading 3'
    elif level == 4:
        paragraph.style = 'Heading 4'
    elif level == 5:
        paragraph.style = 'Heading 5'
    elif level == 6:
        paragraph.style = 'Heading 6'

def process_html_element(doc, element, list_level=0):
    """Process an HTML element and add it to the document."""
    if element.name in ['h1', 'h2', 'h3', 'h4', 'h5', 'h6']:
        level = int(element.name[1])
        heading = doc.add_paragraph(element.get_text())
        apply_heading_style(heading, level)
        return
    
    if element.name == 'p':
        p = doc.add_paragraph()
        process_inline_elements(p, element)
        return
    
    if element.name == 'pre':
        code_block = element.find('code')
        if code_block:
            code_text = code_block.get_text()
            p = doc.add_paragraph(code_text)
            p.style = 'Code Block'
        return
    
    if element.name == 'ul':
        process_list(doc, element, 'List Bullet', list_level)
        return
    
    if element.name == 'ol':
        process_list(doc, element, 'List Number', list_level)
        return
    
    if element.name == 'blockquote':
        p = doc.add_paragraph()
        process_inline_elements(p, element)
        p.style = 'Quote'
        return
    
    if element.name == 'table':
        process_table(doc, element)
        return

def process_inline_elements(paragraph, element):
    """Process inline elements within a paragraph."""
    for child in element.children:
        if isinstance(child, str):
            paragraph.add_run(child)
        elif child.name == 'strong' or child.name == 'b':
            paragraph.add_run(child.get_text()).bold = True
        elif child.name == 'em' or child.name == 'i':
            paragraph.add_run(child.get_text()).italic = True
        elif child.name == 'code':
            run = paragraph.add_run(child.get_text())
            run.font.name = 'Courier New'
            run.font.color.rgb = RGBColor(169, 169, 169)  # Dark gray
        elif child.name == 'a':
            # Add hyperlink
            text = child.get_text()
            url = child.get('href', '')
            run = paragraph.add_run(text)
            run.font.color.rgb = RGBColor(0, 0, 255)  # Blue
            run.font.underline = True
        elif child.name == 'br':
            paragraph.add_run('\n')
        elif child.name:
            # Recursively process other inline elements
            process_inline_elements(paragraph, child)

def process_list(doc, list_element, style_name, level=0):
    """Process a list element (ul or ol)."""
    for li in list_element.find_all('li', recursive=False):
        # Create paragraph with appropriate list style
        p = doc.add_paragraph()
        p.style = style_name
        p.paragraph_format.left_indent = Inches(0.25 * level)
        
        # Process the content of the list item
        for child in li.children:
            if isinstance(child, str):
                p.add_run(child)
            elif child.name in ['ul', 'ol']:
                # Handle nested lists
                next_style = 'List Bullet' if child.name == 'ul' else 'List Number'
                process_list(doc, child, next_style, level + 1)
            else:
                # Process other inline elements
                process_inline_elements(p, child)

def process_table(doc, table_element):
    """Process a table element."""
    # Count rows and columns
    rows = table_element.find_all('tr')
    if not rows:
        return
    
    # Determine the maximum number of cells in any row
    max_cols = 0
    for row in rows:
        cells = row.find_all(['th', 'td'])
        max_cols = max(max_cols, len(cells))
    
    if max_cols == 0:
        return
    
    # Create the table
    table = doc.add_table(rows=len(rows), cols=max_cols)
    table.style = 'Table Grid'
    
    # Fill the table
    for i, row in enumerate(rows):
        cells = row.find_all(['th', 'td'])
        for j, cell in enumerate(cells):
            if j < max_cols:  # Ensure we don't exceed the table dimensions
                # Get the text and apply formatting
                text = cell.get_text().strip()
                table.cell(i, j).text = text
                
                # Apply bold formatting to header cells
                if cell.name == 'th' or i == 0:
                    for paragraph in table.cell(i, j).paragraphs:
                        for run in paragraph.runs:
                            run.bold = True

def parse_xml(xml_string):
    """Parse XML string to an element."""
    return docx.oxml.parse_xml(xml_string)

def add_title_page(doc, title, author, date=None):
    """Add a title page to the document."""
    # Add a section break for the title page
    section = doc.sections[0]
    
    # Set the title page margins
    section.top_margin = Inches(2)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    
    # Add the title
    title_para = doc.add_paragraph()
    title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_run = title_para.add_run(title)
    title_run.bold = True
    title_run.font.size = Pt(24)
    title_run.font.name = 'Times New Roman'
    
    # Add some space
    doc.add_paragraph()
    doc.add_paragraph()
    doc.add_paragraph()
    
    # Add the author
    author_para = doc.add_paragraph()
    author_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    author_run = author_para.add_run(f"By\n{author}")
    author_run.font.size = Pt(16)
    author_run.font.name = 'Times New Roman'
    
    # Add some space
    doc.add_paragraph()
    doc.add_paragraph()
    doc.add_paragraph()
    
    # Add the date
    if date is None:
        date = datetime.now().strftime("%B %d, %Y")
    date_para = doc.add_paragraph()
    date_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    date_run = date_para.add_run(date)
    date_run.font.size = Pt(12)
    date_run.font.name = 'Times New Roman'
    
    # Add page break after title page
    doc.add_page_break()

def add_table_of_contents(doc):
    """Add a table of contents to the document."""
    # Add heading for TOC
    toc_heading = doc.add_paragraph("Table of Contents")
    toc_heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
    toc_heading.style = 'Heading 1'
    
    # Add the TOC field
    paragraph = doc.add_paragraph()
    run = paragraph.add_run()
    
    fldChar1 = create_element('w:fldChar')
    create_attribute(fldChar1, 'w:fldCharType', 'begin')
    
    instrText = create_element('w:instrText')
    create_attribute(instrText, 'xml:space', 'preserve')
    instrText.text = 'TOC \\o "1-3" \\h \\z \\u'  # Include heading levels 1-3
    
    fldChar2 = create_element('w:fldChar')
    create_attribute(fldChar2, 'w:fldCharType', 'separate')
    
    fldChar3 = create_element('w:fldChar')
    create_attribute(fldChar3, 'w:fldCharType', 'end')
    
    run._element.append(fldChar1)
    run._element.append(instrText)
    run._element.append(fldChar2)
    run._element.append(fldChar3)
    
    # Add page break after TOC
    doc.add_page_break()

def add_footer_with_page_numbers(doc):
    """Add a footer with page numbers to the document."""
    # Add footer to all sections except the first one (title page)
    for i, section in enumerate(doc.sections):
        if i > 0:  # Skip the title page
            footer = section.footer
            paragraph = footer.paragraphs[0]
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            add_page_number(paragraph)

def convert_markdown_to_docx(input_path, output_file, title="AI Engineering Thesis", author="Author Name", single_file=False):
    """Convert markdown files to a single docx file.
    
    If single_file is True, input_path is treated as a path to a single markdown file.
    Otherwise, input_path is treated as a directory containing multiple markdown files.
    """
    # Create a new Document
    doc = Document()
    
    # Set up document styles
    setup_styles(doc)
    
    # Add title page
    add_title_page(doc, title, author)
    
    # Add a section break for the main content
    section = doc.add_section()
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    
    # Add table of contents
    add_table_of_contents(doc)
    
    if single_file:
        # Process a single markdown file
        print(f"Processing single file: {input_path}")
        
        # Read markdown content
        with open(input_path, 'r', encoding='utf-8') as f:
            md_content = f.read()
        
        # Convert markdown to HTML
        html = markdown.markdown(md_content, extensions=['tables', 'fenced_code', 'nl2br'])
        
        # Parse HTML with BeautifulSoup
        soup = BeautifulSoup(html, 'html.parser')
        
        # Process each top-level element
        for element in soup.find_all(recursive=False):
            process_html_element(doc, element)
    else:
        # Get all markdown files and sort them by section number
        md_files = [f for f in os.listdir(input_path) if f.endswith('.md')]
        md_files.sort(key=extract_section_number)
        
        print(f"Found {len(md_files)} markdown files to process.")
        
        # Process each markdown file
        for i, md_file in enumerate(md_files):
            print(f"Processing file {i+1}/{len(md_files)}: {md_file}")
            file_path = os.path.join(input_path, md_file)
            
            # Read markdown content
            with open(file_path, 'r', encoding='utf-8') as f:
                md_content = f.read()
            
            # Extract title from filename (remove section number and underscores)
            section_title = md_file.split('.')[0]  # Remove extension
            section_title = re.sub(r'^\d+_', '', section_title)  # Remove section number
            section_title = section_title.replace('_', ' ').title()  # Replace underscores with spaces and capitalize
            
            # Add section title
            heading = doc.add_heading(section_title, level=1)
            heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            # Convert markdown to HTML
            html = markdown.markdown(md_content, extensions=['tables', 'fenced_code', 'nl2br'])
            
            # Parse HTML with BeautifulSoup
            soup = BeautifulSoup(html, 'html.parser')
            
            # Process each top-level element
            for element in soup.find_all(recursive=False):
                process_html_element(doc, element)
            
            # Add page break between sections (except for the last one)
            if i < len(md_files) - 1:
                doc.add_page_break()
    
    # Add footer with page numbers
    add_footer_with_page_numbers(doc)
    
    # Save the document
    doc.save(output_file)
    print(f"Document saved as {output_file}")

if __name__ == "__main__":
    # Set up command-line argument parsing
    parser = argparse.ArgumentParser(description='Convert markdown files to a single Word document')
    parser.add_argument('--input', type=str, help='Input directory containing markdown files',
                        default=r"C:\Users\roger\OneDrive\Desktop\Gauntlet_Projects\PROJECT-REVIEWS\()_SELF-ANALYSIS\AI Engineering Thesis\sections")
    parser.add_argument('--single-file', type=str, help='Path to a single markdown file to convert',
                        default=None)
    parser.add_argument('--output', type=str, help='Output docx file path',
                        default=r"C:\Users\roger\OneDrive\Desktop\Gauntlet_Projects\PROJECT-REVIEWS\()_SELF-ANALYSIS\AI Engineering Thesis\combined_thesis.docx")
    parser.add_argument('--title', type=str, help='Document title',
                        default="AI Engineering Thesis: A Framework for Building with AI")
    parser.add_argument('--author', type=str, help='Author name',
                        default="Roger Pincombe")
    
    args = parser.parse_args()
    
    # Create output directory if it doesn't exist
    os.makedirs(os.path.dirname(args.output), exist_ok=True)
    
    # Convert markdown to docx
    if args.single_file:
        convert_markdown_to_docx(
            args.single_file, 
            args.output,
            title=args.title,
            author=args.author,
            single_file=True
        )
    else:
        convert_markdown_to_docx(
            args.input, 
            args.output,
            title=args.title,
            author=args.author
        ) 